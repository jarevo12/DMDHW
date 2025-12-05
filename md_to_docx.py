#!/usr/bin/env python3
"""
Convert Markdown to Word Document with proper formatting
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set hyperlink style
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def parse_inline_formatting(text, run):
    """Parse and apply inline formatting like bold, italic, code."""
    # Handle bold (**text** or __text__)
    bold_pattern = r'\*\*(.+?)\*\*|__(.+?)__'
    # Handle italic (*text* or _text_)
    italic_pattern = r'\*(.+?)\*|_(.+?)_'
    # Handle inline code (`code`)
    code_pattern = r'`(.+?)`'

    # Simple approach: detect the primary formatting
    if re.match(r'^\*\*.*\*\*$', text) or re.match(r'^__.*__$', text):
        run.bold = True
        text = re.sub(r'^\*\*|\*\*$|^__|__$', '', text)
    elif re.match(r'^\*.*\*$', text) or re.match(r'^_.*_$', text):
        run.italic = True
        text = re.sub(r'^\*|\*$|^_|_$', '', text)
    elif re.match(r'^`.*`$', text):
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        text = re.sub(r'^`|`$', '', text)

    run.text = text

def add_formatted_text(paragraph, text):
    """Add text with inline formatting to a paragraph."""
    # Split text by formatting markers while preserving them
    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\*[^*]+\*|_[^_]+_)', text)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()

        if re.match(r'^\*\*.*\*\*$|^__.*__$', part):
            run.bold = True
            run.text = re.sub(r'^\*\*|\*\*$|^__|__$', '', part)
        elif re.match(r'^`.*`$', part):
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.text = re.sub(r'^`|`$', '', part)
        elif re.match(r'^\*.*\*$|^_.*_$', part):
            run.italic = True
            run.text = re.sub(r'^\*|\*$|^_|_$', '', part)
        else:
            run.text = part

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    p = doc.add_paragraph()
                    for code_line in code_lines:
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle headers
        if line.startswith('#'):
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)

                # Remove markdown links from headers
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

                heading = doc.add_heading(text, level=level)
                if level == 1:
                    heading.runs[0].font.size = Pt(20)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                elif level == 2:
                    heading.runs[0].font.size = Pt(16)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                elif level == 3:
                    heading.runs[0].font.size = Pt(14)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        # Handle horizontal rules
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 50)

        # Handle tables
        elif '|' in line and not in_table:
            # Start collecting table rows
            table_rows = [line]
            i += 1
            # Collect all table rows
            while i < len(lines) and '|' in lines[i]:
                table_rows.append(lines[i].rstrip())
                i += 1

            # Parse and create table
            if len(table_rows) >= 2:  # Header + separator at minimum
                # Parse header
                headers = [cell.strip() for cell in table_rows[0].split('|')[1:-1]]

                # Parse data rows (skip separator row)
                data_rows = []
                for row in table_rows[2:]:
                    cells = [cell.strip() for cell in row.split('|')[1:-1]]
                    if cells:
                        data_rows.append(cells)

                # Create table
                if data_rows:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
                    table.style = 'Light Grid Accent 1'

                    # Add headers
                    for j, header in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = header
                        # Bold header text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                    # Add data
                    for row_idx, row_data in enumerate(data_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(headers):
                                # Handle checkmarks and symbols
                                cell_data = cell_data.replace('✅', '✓')
                                cell_data = cell_data.replace('❌', '✗')
                                cell_data = cell_data.replace('⚠️', '⚠')
                                table.rows[row_idx + 1].cells[col_idx].text = cell_data

                doc.add_paragraph()  # Add spacing after table
            continue

        # Handle unordered lists
        elif re.match(r'^[\*\-\+]\s+', line):
            text = re.sub(r'^[\*\-\+]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)

        # Handle ordered lists
        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)

        # Handle blockquotes
        elif line.startswith('>'):
            text = re.sub(r'^>\s*', '', line)
            p = doc.add_paragraph()
            p.style = 'Intense Quote'
            add_formatted_text(p, text)

        # Handle empty lines
        elif not line.strip():
            if i > 0:  # Don't add space at the beginning
                doc.add_paragraph()

        # Handle regular paragraphs
        else:
            if line.strip():
                p = doc.add_paragraph()
                add_formatted_text(p, line)

        i += 1

    # Save the document
    doc.save(docx_file)
    print(f"✓ Successfully created {docx_file}")

if __name__ == '__main__':
    md_file = '/workspaces/DMDHW/habit-tracker-analysis.md'
    docx_file = '/workspaces/DMDHW/habit-tracker-analysis.docx'

    convert_md_to_docx(md_file, docx_file)
