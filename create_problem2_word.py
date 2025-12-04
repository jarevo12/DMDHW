#!/usr/bin/env python3
"""Convert Problem 2 Solution markdown to Word document."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_formatted_text(paragraph, text, bold=False, italic=False, code=False):
    """Add formatted text to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    return run

def process_line(doc, line, in_code_block=False):
    """Process a single line of markdown and add to document."""

    # Skip horizontal rules
    if line.strip() in ['---', '***', '___']:
        doc.add_paragraph()
        return in_code_block

    # Code blocks
    if line.strip().startswith('```'):
        return not in_code_block

    if in_code_block:
        p = doc.add_paragraph(line, style='Normal')
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        return in_code_block

    # Headings
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)

    # Lists
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        p = doc.add_paragraph(text, style='List Bullet')

    # Numbered lists
    elif re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        p = doc.add_paragraph(text, style='List Number')

    # Empty lines
    elif not line.strip():
        pass  # Skip empty lines

    # Regular paragraphs
    elif not line.startswith('#'):
        # Parse inline formatting
        p = doc.add_paragraph()

        # Simple bold/italic parsing
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', line)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                add_formatted_text(p, part[2:-2], bold=True)
            elif part.startswith('*') and part.endswith('*'):
                add_formatted_text(p, part[1:-1], italic=True)
            elif part.startswith('`') and part.endswith('`'):
                add_formatted_text(p, part[1:-1], code=True)
            elif part:
                p.add_run(part)

    return in_code_block

def create_word_document(markdown_file, output_file):
    """Create Word document from markdown file."""

    # Create document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Process each line
    in_code_block = False
    for line in lines:
        in_code_block = process_line(doc, line.rstrip(), in_code_block)

    # Save document
    doc.save(output_file)
    print(f"Word document created: {output_file}")

if __name__ == "__main__":
    markdown_file = "/workspaces/DMDHW/Problem2_Solution.md"
    output_file = "/workspaces/DMDHW/Problem2_Solution.docx"

    create_word_document(markdown_file, output_file)
